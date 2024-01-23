from docx import Document

# Create a new Word document
doc = Document()

# Add content to the document
doc.add_heading('Document Title', level=1)
doc.add_paragraph('This is a sample document created with python-docx.')

# Save the document
doc.save('example.docx')
